#!/usr/bin/env python3
"""ats_docx.py — convert resume Markdown into ATS-safe .docx / .txt (and optional .pdf).

This script is **standalone and dependency-light** on purpose: it is injected into an
ephemeral Hermes sandbox container (see ``services/core/hermes/render.py``) and executed
there with no network access. Its only hard dependency is ``python-docx``; LibreOffice
(``soffice``) is used for PDF when present and is skipped gracefully when it is not.

Usage
-----
    python ats_docx.py --input resume.md --outdir /work --basename resume
    python ats_docx.py --input resume.md --outdir /work --basename resume --no-pdf
    python ats_docx.py --input resume.md --outdir /work --basename resume --json-out result.json

On success a single JSON object is printed to stdout::

    {"ok": true, "docx": "...", "txt": "...", "pdf": "..."|null, "pdf_error": null,
     "warnings": [...], "stats": {...}}

ATS-safe rules enforced (from the Hermes build contract)
--------------------------------------------------------
* single column — one document section, no column definitions, no side panels
* no tables (Markdown tables are flattened into plain text lines)
* no text boxes, no shapes, no images (nothing that can emit one is ever called)
* no headers or footers (never populated; page numbers are deliberately omitted)
* standard fonts only — Calibri by default, Arial/Liberation as fallbacks, 10-12pt
* real heading paragraphs (built-in Heading 1/2 styles, restyled to plain black bold)
  so parsers get true outline levels instead of "bold text that looks like a heading"
* bullets are the literal character "•" plus a hanging indent — no numbering
  definitions, which some parsers drop entirely
* contact details are plain text lines (no hyperlink fields, no icons, no glyphs)
* smart quotes, dashes, arrows and emoji are normalised to ASCII

Honesty note: these rules follow widely published ATS parsing guidance. They make the
document reliably *machine-readable*; no formatting choice can guarantee a particular
score in any specific vendor's parser.
"""

from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import subprocess
import sys
import unicodedata
from dataclasses import dataclass, field
from typing import Dict, List, Optional, Tuple

# --------------------------------------------------------------------------------------
# Constants
# --------------------------------------------------------------------------------------

#: Canonical section headings, in the order the ResumeArchitect emits them.
CANONICAL_SECTIONS: Tuple[str, ...] = (
    "PROFESSIONAL SUMMARY",
    "CORE COMPETENCIES",
    "PROFESSIONAL EXPERIENCE",
    "PROJECTS",
    "EDUCATION",
    "CERTIFICATIONS",
    "TECHNICAL SKILLS",
)

#: Common synonyms mapped onto the canonical headings, so a slightly-off input still
#: produces the canonical wording in the rendered document.
SECTION_ALIASES: Dict[str, str] = {
    "SUMMARY": "PROFESSIONAL SUMMARY",
    "PROFILE": "PROFESSIONAL SUMMARY",
    "PROFESSIONAL PROFILE": "PROFESSIONAL SUMMARY",
    "EXECUTIVE SUMMARY": "PROFESSIONAL SUMMARY",
    "OBJECTIVE": "PROFESSIONAL SUMMARY",
    "CORE COMPETENCE": "CORE COMPETENCIES",
    "COMPETENCIES": "CORE COMPETENCIES",
    "KEY SKILLS": "CORE COMPETENCIES",
    "AREAS OF EXPERTISE": "CORE COMPETENCIES",
    "EXPERIENCE": "PROFESSIONAL EXPERIENCE",
    "WORK EXPERIENCE": "PROFESSIONAL EXPERIENCE",
    "EMPLOYMENT HISTORY": "PROFESSIONAL EXPERIENCE",
    "PROFESSIONAL BACKGROUND": "PROFESSIONAL EXPERIENCE",
    "SELECTED PROJECTS": "PROJECTS",
    "KEY PROJECTS": "PROJECTS",
    "EDUCATION & TRAINING": "EDUCATION",
    "ACADEMIC BACKGROUND": "EDUCATION",
    "CERTIFICATION": "CERTIFICATIONS",
    "CERTIFICATIONS & LICENSES": "CERTIFICATIONS",
    "LICENSES & CERTIFICATIONS": "CERTIFICATIONS",
    "SKILLS": "TECHNICAL SKILLS",
    "TECHNICAL EXPERTISE": "TECHNICAL SKILLS",
    "TOOLS & TECHNOLOGIES": "TECHNICAL SKILLS",
}

BULLET_CHAR = "•"  # •
TXT_BULLET = "- "  # plain ASCII in the .txt flavour (safest for copy/paste forms)

#: Fonts we allow. First entry is the default; all are metrically standard and present
#: (or substitutable via fonts-liberation) in the sandbox image.
SAFE_FONTS: Tuple[str, ...] = ("Calibri", "Arial", "Helvetica", "Times New Roman", "Georgia", "Verdana")

#: Glyph normalisation: everything a language model tends to emit -> ASCII.
GLYPH_MAP: Dict[str, str] = {
    "‘": "'", "’": "'", "‚": "'", "‛": "'",
    "“": '"', "”": '"', "„": '"', "‟": '"',
    "′": "'", "″": '"', "´": "'", "`": "'",
    "–": "-", "—": "-", "―": "-", "−": "-", "⁃": "-",
    "…": "...",
    " ": " ", " ": " ", " ": " ", " ": " ", " ": " ",
    " ": " ", " ": " ", " ": " ",
    "​": "", "‌": "", "‍": "", "﻿": "",
    "←": "<-", "→": "->", "↔": "<->", "⇒": "=>",
    "•": BULLET_CHAR, "‣": BULLET_CHAR, "▪": BULLET_CHAR,
    "●": BULLET_CHAR, "◦": BULLET_CHAR, "·": BULLET_CHAR,
    "∙": BULLET_CHAR, "■": BULLET_CHAR, "❖": BULLET_CHAR,
    "✓": "", "✔": "", "✗": "",
    "×": "x", "⁄": "/", "­": "",
    "™": "(TM)", "®": "(R)",
}

# Regexes
RE_HEADING = re.compile(r"^(#{1,6})\s*(.+?)\s*#*\s*$")
RE_BULLET = re.compile(r"^(\s*)(?:[-*+•]|\d{1,2}[.)])\s+(.*)$")
RE_HR = re.compile(r"^\s*(?:-{3,}|\*{3,}|_{3,})\s*$")
RE_SETEXT_H1 = re.compile(r"^\s*={3,}\s*$")
RE_TABLE_SEP = re.compile(r"^\s*\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)*\|?\s*$")
RE_HTML_COMMENT = re.compile(r"<!--.*?-->", re.DOTALL)
RE_HTML_TAG = re.compile(r"</?[A-Za-z][^>]*>")
RE_LINK = re.compile(r"\[([^\]]*)\]\(([^)\s]+)(?:\s+\"[^\"]*\")?\)")
RE_IMAGE = re.compile(r"!\[([^\]]*)\]\(([^)\s]+)(?:\s+\"[^\"]*\")?\)")
RE_EMAIL = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
RE_PHONE = re.compile(r"(?:\+?\d[\d\-\s().]{7,}\d)")
RE_URLISH = re.compile(r"(?:https?://|www\.)\S+|linkedin\.com/\S+|github\.com/\S+", re.IGNORECASE)


# --------------------------------------------------------------------------------------
# Document model
# --------------------------------------------------------------------------------------


@dataclass
class Block:
    """One rendered paragraph."""

    kind: str  # name | contact | section | subheading | bullet | para
    text: str
    level: int = 0  # bullet nesting depth (0 or 1)


@dataclass
class ParsedResume:
    name: str = ""
    contact: List[str] = field(default_factory=list)
    blocks: List[Block] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)
    sections_found: List[str] = field(default_factory=list)


# --------------------------------------------------------------------------------------
# Text sanitising
# --------------------------------------------------------------------------------------


def sanitize_text(text: str) -> str:
    """Normalise unicode punctuation to ASCII and strip decorative/invisible glyphs.

    Letters (including accented ones) and ordinary punctuation are preserved; emoji,
    dingbats, private-use characters and control/format codepoints are dropped, because
    they routinely corrupt resume text extraction.
    """
    text = unicodedata.normalize("NFKC", text)
    out_chars: List[str] = []
    for ch in text:
        if ch in GLYPH_MAP:
            out_chars.append(GLYPH_MAP[ch])
            continue
        if ch in ("\n", "\t"):
            out_chars.append(ch)
            continue
        category = unicodedata.category(ch)
        if category in ("Cc", "Cf", "Co", "Cs", "Cn"):
            continue  # control / format / private use / surrogate / unassigned
        if category in ("So", "Sk") and ch != BULLET_CHAR:
            continue  # emoji, dingbats, modifier symbols
        out_chars.append(ch)
    cleaned = "".join(out_chars)
    cleaned = cleaned.replace("\t", "    ")
    # Collapse runs of spaces (but never touch newlines / leading indentation semantics
    # are handled by the parser before this point).
    return cleaned


def strip_inline_markup(text: str) -> str:
    """Remove Markdown decoration, keeping the words a parser needs."""
    text = RE_IMAGE.sub(r"\1", text)  # images become their alt text (never rendered)

    def _link(match: "re.Match[str]") -> str:
        label, url = match.group(1).strip(), match.group(2).strip()
        if not label:
            return url
        bare = url.replace("https://", "").replace("http://", "").rstrip("/")
        if label.rstrip("/") in (url.rstrip("/"), bare) or url.startswith("mailto:"):
            return label
        return f"{label} ({bare})"

    text = RE_LINK.sub(_link, text)
    text = RE_HTML_TAG.sub("", text)
    text = text.replace("`", "")
    text = re.sub(r"~~(.+?)~~", r"\1", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def parse_runs(text: str) -> List[Tuple[str, bool, bool]]:
    """Split a line into ``(text, bold, italic)`` runs from ``**bold**`` / ``*italic*``.

    Bold/italic are the only formatting we keep: they are safe for every parser and help
    a human reader scan job titles and company names.
    """
    tokens: List[Tuple[str, bool, bool]] = []
    pattern = re.compile(r"(\*\*\*|___|\*\*|__|\*|_)")
    bold = False
    italic = False
    pos = 0
    buffer: List[str] = []

    def flush() -> None:
        if buffer:
            tokens.append(("".join(buffer), bold, italic))
            buffer.clear()

    while pos < len(text):
        match = pattern.search(text, pos)
        if not match:
            buffer.append(text[pos:])
            break
        # `_` inside a word (snake_case, file_name) is literal, not emphasis.
        marker = match.group(1)
        start = match.start()
        if marker in ("_", "__", "___"):
            prev_ch = text[start - 1] if start > 0 else " "
            next_idx = match.end()
            next_ch = text[next_idx] if next_idx < len(text) else " "
            if (prev_ch.isalnum() and next_ch.isalnum()) or (prev_ch.isalnum() and marker == "_"):
                buffer.append(text[pos : match.end()])
                pos = match.end()
                continue
        buffer.append(text[pos:start])
        flush()
        if marker in ("***", "___"):
            bold = not bold
            italic = not italic
        elif marker in ("**", "__"):
            bold = not bold
        else:
            italic = not italic
        pos = match.end()
    flush()
    return [(t, b, i) for (t, b, i) in tokens if t]


def canonicalise_section(raw: str) -> Optional[str]:
    """Return the canonical heading name for ``raw``, or None if it is not a section."""
    probe = re.sub(r"[^A-Z& ]+", " ", raw.upper())
    probe = re.sub(r"\s+", " ", probe).strip()
    if not probe:
        return None
    if probe in CANONICAL_SECTIONS:
        return probe
    if probe in SECTION_ALIASES:
        return SECTION_ALIASES[probe]
    compact = probe.replace("AND", "&").replace(" ", "")
    for canonical in CANONICAL_SECTIONS:
        if compact == canonical.replace(" ", ""):
            return canonical
    for alias, canonical in SECTION_ALIASES.items():
        if compact == alias.replace("AND", "&").replace(" ", ""):
            return canonical
    return None


def looks_like_contact(line: str) -> bool:
    """Heuristic: does this pre-section line carry contact details?"""
    return bool(
        RE_EMAIL.search(line)
        or RE_URLISH.search(line)
        or RE_PHONE.search(line)
        or "|" in line
        or re.search(r"\b(?:phone|mobile|email|linkedin|github|portfolio)\b", line, re.IGNORECASE)
    )


# --------------------------------------------------------------------------------------
# Markdown parsing
# --------------------------------------------------------------------------------------


def parse_markdown(markdown: str) -> ParsedResume:
    """Parse resume Markdown into a flat, single-column block list."""
    doc = ParsedResume()
    text = RE_HTML_COMMENT.sub("", markdown)
    text = sanitize_text(text.replace("\r\n", "\n").replace("\r", "\n"))

    # Drop fenced code blocks entirely (they are never legitimate resume content and
    # their monospace styling confuses parsers).
    lines: List[str] = []
    in_fence = False
    for raw_line in text.split("\n"):
        if raw_line.strip().startswith("```") or raw_line.strip().startswith("~~~"):
            in_fence = not in_fence
            if not in_fence:
                doc.warnings.append("A fenced code block was removed from the resume body.")
            continue
        if in_fence:
            continue
        lines.append(raw_line)

    seen_section = False
    pending_para: List[str] = []

    def flush_para() -> None:
        if pending_para:
            joined = strip_inline_markup(" ".join(part.strip() for part in pending_para).strip())
            if joined:
                doc.blocks.append(Block("para", joined))
            pending_para.clear()

    index = 0
    total = len(lines)
    while index < total:
        raw_line = lines[index]
        line = raw_line.rstrip()
        stripped = line.strip()
        index += 1

        if not stripped:
            flush_para()
            continue

        if RE_HR.match(stripped):
            flush_para()
            continue  # horizontal rules carry no content

        # Setext H1 ("Name" followed by "====") -> promote the previous paragraph.
        if RE_SETEXT_H1.match(stripped) and doc.blocks and doc.blocks[-1].kind == "para":
            promoted = doc.blocks.pop()
            if not doc.name:
                doc.name = promoted.text
            else:
                doc.blocks.append(Block("subheading", promoted.text))
            continue

        # Markdown tables are forbidden in the output: flatten each row to plain text.
        if stripped.startswith("|") or RE_TABLE_SEP.match(stripped):
            if RE_TABLE_SEP.match(stripped):
                continue  # the |---|---| separator row carries no content
            cells = [strip_inline_markup(cell.strip()) for cell in stripped.strip("|").split("|")]
            cells = [cell for cell in cells if cell]
            if cells:
                flush_para()
                doc.blocks.append(Block("bullet", " - ".join(cells), 0))
                if "table" not in " ".join(doc.warnings):
                    doc.warnings.append("A Markdown table was flattened into plain lines (tables are not ATS-safe).")
            continue

        heading_match = RE_HEADING.match(stripped)
        if heading_match:
            flush_para()
            hashes, heading_text = heading_match.group(1), strip_inline_markup(heading_match.group(2))
            canonical = canonicalise_section(heading_text)
            if canonical:
                seen_section = True
                doc.blocks.append(Block("section", canonical))
                if canonical not in doc.sections_found:
                    doc.sections_found.append(canonical)
            elif not seen_section and len(hashes) == 1 and not doc.name:
                doc.name = heading_text
            elif len(hashes) <= 2 and heading_text.upper() == heading_text and len(heading_text) > 2:
                # An unrecognised but clearly section-like heading (e.g. "AWARDS").
                seen_section = True
                doc.blocks.append(Block("section", heading_text))
            else:
                doc.blocks.append(Block("subheading", heading_text))
            continue

        bullet_match = RE_BULLET.match(raw_line)
        if bullet_match:
            flush_para()
            indent = len(bullet_match.group(1).expandtabs(4))
            level = 1 if indent >= 2 else 0
            content = strip_inline_markup(bullet_match.group(2))
            if content:
                doc.blocks.append(Block("bullet", content, level))
            continue

        # Plain line.
        if not seen_section:
            plain = strip_inline_markup(stripped)
            if not doc.name and not looks_like_contact(plain):
                doc.name = plain
                continue
            if plain:
                doc.contact.append(re.sub(r"\s*\|\s*", " | ", plain))
            continue

        # Inside a section: a short bold-only line is a company/role/date subheading.
        if re.fullmatch(r"\*\*.+\*\*[:.]?", stripped) or re.fullmatch(r"__.+__[:.]?", stripped):
            flush_para()
            doc.blocks.append(Block("subheading", strip_inline_markup(stripped)))
            continue

        pending_para.append(stripped)

    flush_para()

    if not doc.name:
        doc.warnings.append("No candidate name was found at the top of the resume.")
    if not doc.contact:
        doc.warnings.append("No contact block detected (email / phone / location / LinkedIn URL).")
    missing = [section for section in CANONICAL_SECTIONS if section not in doc.sections_found]
    if missing:
        doc.warnings.append("Canonical sections absent from the input: " + ", ".join(missing))

    # Drop section headings with no content beneath them (an empty "PROJECTS" heading
    # is worse than no heading at all for both parsers and readers).
    pruned: List[Block] = []
    for position, block in enumerate(doc.blocks):
        if block.kind == "section":
            following = doc.blocks[position + 1]  if position + 1 < len(doc.blocks) else None
            if following is None or following.kind == "section":
                doc.warnings.append(f"Dropped empty section '{block.text}'.")
                continue
        pruned.append(block)
    doc.blocks = pruned
    if doc.sections_found:
        expected = [s for s in CANONICAL_SECTIONS if s in doc.sections_found]
        actual = [s for s in doc.sections_found if s in CANONICAL_SECTIONS]
        if expected != actual:
            doc.warnings.append(
                "Canonical sections are not in the contract order "
                f"({' > '.join(actual)}); rendered in input order."
            )
    return doc


# --------------------------------------------------------------------------------------
# DOCX rendering
# --------------------------------------------------------------------------------------


def _clear_theme_fonts(style) -> None:
    """Remove theme font references so our explicit font actually wins in Word."""
    try:
        from docx.oxml.ns import qn

        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            return
        for attribute in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
            rfonts.attrib.pop(qn("w:" + attribute), None)
    except Exception:
        pass  # cosmetic only


def _force_font(style, font_name: str, size_pt: float, *, bold: Optional[bool] = None) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    font = style.font
    font.name = font_name
    font.size = Pt(size_pt)
    font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    font.all_caps = False
    font.small_caps = False
    if bold is not None:
        font.bold = bold
    _clear_theme_fonts(style)
    try:
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), font_name)
        rfonts.set(qn("w:cs"), font_name)
    except Exception:
        pass


def build_docx(doc_model: ParsedResume, out_path: str, font_name: str, size_pt: float) -> Dict[str, object]:
    """Write the ATS-safe .docx. Returns rendering stats."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.shared import Inches, Pt

    document = Document()

    # --- page geometry: one section, single column, generous margins -----------------
    section = document.sections[0]
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.different_first_page_header_footer = False
    # Headers/footers are never populated. Explicitly blank anything the default
    # template shipped so no stray content reaches the parser.
    header_footer_parts = [getattr(section, attr, None) for attr in
                           ("header", "footer", "first_page_header", "first_page_footer",
                            "even_page_header", "even_page_footer")]
    for part in header_footer_parts:
        if part is None:
            continue
        try:
            for paragraph in part.paragraphs:
                for run in list(paragraph.runs):
                    run.text = ""
        except Exception:
            pass

    # --- styles ----------------------------------------------------------------------
    normal = document.styles["Normal"]
    _force_font(normal, font_name, size_pt, bold=False)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.widow_control = True
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Real heading styles (restyled to plain black bold — no theme blue, no borders).
    heading1 = document.styles["Heading 1"]
    _force_font(heading1, font_name, min(size_pt + 5, 18), bold=True)
    heading1.paragraph_format.space_before = Pt(0)
    heading1.paragraph_format.space_after = Pt(2)
    heading1.paragraph_format.keep_with_next = True

    heading2 = document.styles["Heading 2"]
    _force_font(heading2, font_name, min(size_pt + 1.5, 14), bold=True)
    heading2.paragraph_format.space_before = Pt(10)
    heading2.paragraph_format.space_after = Pt(3)
    heading2.paragraph_format.keep_with_next = True

    stats = {"bullets": 0, "paragraphs": 0, "sections": 0, "words": 0}

    def add_runs(paragraph, text: str, *, bold_all: bool = False) -> None:
        for chunk, bold, italic in parse_runs(text) or [(text, False, False)]:
            run = paragraph.add_run(chunk)
            run.bold = bool(bold or bold_all)
            run.italic = bool(italic)
            run.font.name = font_name
        stats["words"] += len(text.split())

    # --- name + contact ---------------------------------------------------------------
    if doc_model.name:
        paragraph = document.add_paragraph(style="Heading 1")
        add_runs(paragraph, doc_model.name, bold_all=True)
    for contact_line in doc_model.contact:
        paragraph = document.add_paragraph(style="Normal")
        paragraph.paragraph_format.space_after = Pt(0)
        add_runs(paragraph, contact_line)
        stats["paragraphs"] += 1

    # --- body -------------------------------------------------------------------------
    for block in doc_model.blocks:
        if block.kind == "section":
            paragraph = document.add_paragraph(style="Heading 2")
            add_runs(paragraph, block.text.upper(), bold_all=True)
            stats["sections"] += 1
        elif block.kind == "subheading":
            paragraph = document.add_paragraph(style="Normal")
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(1)
            paragraph.paragraph_format.keep_with_next = True
            add_runs(paragraph, block.text, bold_all=True)
            stats["paragraphs"] += 1
        elif block.kind == "bullet":
            paragraph = document.add_paragraph(style="Normal")
            indent = 0.22 + (0.22 * min(block.level, 1))
            paragraph.paragraph_format.left_indent = Inches(indent)
            paragraph.paragraph_format.first_line_indent = Inches(-0.18)
            paragraph.paragraph_format.space_after = Pt(2)
            # Literal bullet character + hanging indent: no numbering definition, which
            # some ATS parsers silently discard along with the text.
            run = paragraph.add_run(BULLET_CHAR + " ")
            run.font.name = font_name
            add_runs(paragraph, block.text)
            stats["bullets"] += 1
        else:  # para
            paragraph = document.add_paragraph(style="Normal")
            paragraph.paragraph_format.space_after = Pt(4)
            add_runs(paragraph, block.text)
            stats["paragraphs"] += 1

    core = document.core_properties
    core.title = doc_model.name or "Resume"
    core.author = doc_model.name or "Resume"
    core.comments = "Generated by Hermes - ATS-safe single-column resume."

    document.save(out_path)
    stats["word_count"] = stats.pop("words")
    stats["page_estimate"] = max(1, round(stats["word_count"] / 550.0 + 0.35))
    return stats


# --------------------------------------------------------------------------------------
# TXT rendering
# --------------------------------------------------------------------------------------


def build_txt(doc_model: ParsedResume, out_path: str) -> int:
    """Write the plain-text flavour (for paste-into-web-form applications)."""
    lines: List[str] = []
    if doc_model.name:
        lines.append(doc_model.name.upper())
    lines.extend(doc_model.contact)
    for block in doc_model.blocks:
        if block.kind == "section":
            lines.extend(["", block.text.upper(), "=" * len(block.text)])
        elif block.kind == "subheading":
            lines.extend(["", block.text])
        elif block.kind == "bullet":
            lines.append(("    " if block.level else "") + TXT_BULLET + block.text)
        else:
            lines.extend(["", block.text])
    body = "\n".join(lines).strip() + "\n"
    body = re.sub(r"\n{3,}", "\n\n", body)
    with open(out_path, "w", encoding="utf-8", newline="\n") as handle:
        handle.write(body)
    return len(body.split())


# --------------------------------------------------------------------------------------
# PDF rendering (optional)
# --------------------------------------------------------------------------------------


def find_soffice() -> Optional[str]:
    """Locate a LibreOffice binary, or None when PDF support is not installed."""
    for candidate in ("soffice", "libreoffice", "soffice.bin"):
        found = shutil.which(candidate)
        if found:
            return found
    for path in (
        "/usr/lib/libreoffice/program/soffice",
        "/usr/lib/libreoffice/program/soffice.bin",
        "/opt/libreoffice/program/soffice",
    ):
        if os.path.exists(path) and os.access(path, os.X_OK):
            return path
    return None


def build_pdf(docx_path: str, outdir: str, basename: str, timeout: int = 180) -> Tuple[Optional[str], Optional[str]]:
    """Convert the .docx to PDF with LibreOffice. Returns ``(path, error)``.

    Degrades gracefully: a missing binary, a crash, or a timeout yields
    ``(None, reason)`` and never aborts the render.
    """
    soffice = find_soffice()
    if not soffice:
        return None, "LibreOffice (soffice) is not installed in the sandbox image; PDF skipped."

    # LibreOffice insists on a writable user profile. The container root filesystem is
    # read-only and /tmp is a small tmpfs, so the profile goes in the bind-mounted
    # workspace instead.
    profile_dir = os.path.join(outdir, ".lo-profile")
    os.makedirs(profile_dir, exist_ok=True)
    command = [
        soffice,
        f"-env:UserInstallation=file://{profile_dir}",
        "--headless",
        "--norestore",
        "--nolockcheck",
        "--nodefault",
        "--nofirststartwizard",
        "--convert-to",
        "pdf:writer_pdf_Export",
        "--outdir",
        outdir,
        docx_path,
    ]
    env = dict(os.environ)
    env.setdefault("HOME", outdir)
    env["SAL_USE_VCLPLUGIN"] = "svp"  # no X11/GTK backend needed
    try:
        completed = subprocess.run(
            command,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            timeout=timeout,
            env=env,
            check=False,
        )
    except subprocess.TimeoutExpired:
        return None, f"LibreOffice timed out after {timeout}s; PDF skipped."
    except OSError as exc:
        return None, f"Could not launch LibreOffice: {exc}; PDF skipped."

    produced = os.path.join(outdir, f"{basename}.pdf")
    if os.path.exists(produced) and os.path.getsize(produced) > 0:
        return produced, None
    output = (completed.stdout or b"").decode("utf-8", errors="replace").strip()
    tail = output[-600:] if output else "(no output)"
    return None, f"LibreOffice exited {completed.returncode} without producing a PDF: {tail}"


# --------------------------------------------------------------------------------------
# CLI
# --------------------------------------------------------------------------------------


def main(argv: Optional[List[str]] = None) -> int:
    parser = argparse.ArgumentParser(
        prog="ats_docx.py",
        description="Render resume Markdown as an ATS-safe .docx (+ .txt, optional .pdf).",
    )
    parser.add_argument("--input", required=True, help="Path to the resume Markdown file.")
    parser.add_argument("--outdir", default=".", help="Directory to write outputs into (default: cwd).")
    parser.add_argument("--basename", default="resume", help="Output file stem (default: resume).")
    parser.add_argument("--font", default=SAFE_FONTS[0], help=f"Body font, one of {', '.join(SAFE_FONTS)}.")
    parser.add_argument("--font-size", type=float, default=11.0, help="Body font size in points (10-12).")
    parser.add_argument("--no-pdf", action="store_true", help="Skip the LibreOffice PDF step.")
    parser.add_argument("--pdf-timeout", type=int, default=180, help="Seconds allowed for PDF conversion.")
    parser.add_argument("--json-out", default=None, help="Also write the result JSON to this path.")
    args = parser.parse_args(argv)

    result: Dict[str, object] = {
        "ok": False,
        "docx": None,
        "txt": None,
        "pdf": None,
        "pdf_error": None,
        "warnings": [],
        "stats": {},
    }

    def finish(code: int) -> int:
        payload = json.dumps(result, ensure_ascii=False)
        if args.json_out:
            try:
                with open(args.json_out, "w", encoding="utf-8") as handle:
                    handle.write(payload)
            except OSError as exc:  # pragma: no cover
                sys.stderr.write(f"warning: could not write --json-out: {exc}\n")
        sys.stdout.write(payload + "\n")
        sys.stdout.flush()
        return code

    # Validate inputs loudly.
    font_name = args.font if args.font in SAFE_FONTS else SAFE_FONTS[0]
    warnings: List[str] = []
    if font_name != args.font:
        warnings.append(f"Font {args.font!r} is not in the ATS-safe list; using {font_name}.")
    size_pt = args.font_size
    if not 10.0 <= size_pt <= 12.0:
        warnings.append(f"Font size {size_pt} is outside the ATS-safe 10-12pt range; clamped.")
        size_pt = min(max(size_pt, 10.0), 12.0)

    try:
        with open(args.input, "r", encoding="utf-8", errors="replace") as handle:
            markdown = handle.read()
    except OSError as exc:
        result["error"] = f"Cannot read --input {args.input!r}: {exc}"
        result["warnings"] = warnings
        sys.stderr.write(str(result["error"]) + "\n")
        return finish(2)

    if not markdown.strip():
        result["error"] = f"Input file {args.input!r} is empty; nothing to render."
        result["warnings"] = warnings
        sys.stderr.write(str(result["error"]) + "\n")
        return finish(2)

    try:
        import docx  # noqa: F401  (import probe only)
    except Exception as exc:
        result["error"] = (
            f"python-docx is not available in the sandbox image ({exc}). "
            "Rebuild hermes-sandbox: docker compose --profile build-only build hermes-sandbox"
        )
        result["warnings"] = warnings
        sys.stderr.write(str(result["error"]) + "\n")
        return finish(3)

    outdir = os.path.abspath(args.outdir)
    os.makedirs(outdir, exist_ok=True)
    basename = re.sub(r"[^A-Za-z0-9._-]+", "-", args.basename).strip("-.") or "resume"

    model = parse_markdown(markdown)
    warnings.extend(model.warnings)

    docx_path = os.path.join(outdir, f"{basename}.docx")
    txt_path = os.path.join(outdir, f"{basename}.txt")

    try:
        stats = build_docx(model, docx_path, font_name, size_pt)
    except Exception as exc:
        result["error"] = f"DOCX rendering failed: {type(exc).__name__}: {exc}"
        result["warnings"] = warnings
        sys.stderr.write(str(result["error"]) + "\n")
        return finish(4)

    try:
        txt_words = build_txt(model, txt_path)
    except Exception as exc:
        result["error"] = f"TXT rendering failed: {type(exc).__name__}: {exc}"
        result["warnings"] = warnings
        result["docx"] = docx_path
        sys.stderr.write(str(result["error"]) + "\n")
        return finish(5)

    pdf_path: Optional[str] = None
    pdf_error: Optional[str] = None
    if args.no_pdf:
        pdf_error = "PDF generation disabled via --no-pdf."
    else:
        pdf_path, pdf_error = build_pdf(docx_path, outdir, basename, timeout=args.pdf_timeout)
        if pdf_error:
            warnings.append(pdf_error)

    page_estimate = stats.get("page_estimate", 1)
    if isinstance(page_estimate, (int, float)) and page_estimate > 2:
        warnings.append(
            f"Estimated {page_estimate} pages ({stats.get('word_count')} words); most ATS-screened "
            "resumes should fit 1-2 pages."
        )

    result.update(
        {
            "ok": True,
            "docx": docx_path,
            "txt": txt_path,
            "pdf": pdf_path,
            "pdf_error": pdf_error,
            "warnings": warnings,
            "stats": {
                "name": model.name,
                "contact_lines": len(model.contact),
                "sections": model.sections_found,
                "bullets": stats.get("bullets"),
                "paragraphs": stats.get("paragraphs"),
                "word_count": stats.get("word_count"),
                "txt_word_count": txt_words,
                "page_estimate": page_estimate,
                "font": font_name,
                "font_size_pt": size_pt,
                "bullet_char": BULLET_CHAR,
            },
        }
    )
    return finish(0)


if __name__ == "__main__":  # pragma: no cover
    sys.exit(main())
